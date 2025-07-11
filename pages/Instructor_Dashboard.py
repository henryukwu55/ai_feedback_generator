import streamlit as st
st.set_page_config(initial_sidebar_state="expanded")
from utils.session import initialize_session, restrict_access, update_activity
from utils.canvas_integration import fetch_submissions
from utils.grading import fetch_assignment_rubric, get_ai_feedback, save_grading_record, convert_rubric_to_table, format_text_for_display
from database.operations import get_grading_records
from database.connection import get_db_connection
from utils.security import get_all_grading_data
from utils.azure_storage import *
import pandas as pd
import docx
import io
import os
import zipfile
from PyPDF2 import PdfReader
from xml.etree import ElementTree as ET


import streamlit as st

# Remove ONLY the "About" option from the menu
hide_about_option = """
<style>
/* Hide the About menu item */
li[data-testid="main-menu-about"] {
    display: none !important;
}
</style>
"""
st.markdown(hide_about_option, unsafe_allow_html=True)


# Initialize session and security
def track_activity():
    if st.session_state.get("logged_in"):
        update_activity()

# Theme customization
current_theme = st.get_option("theme.base")
if current_theme == "dark":
    st.markdown("""
        <style>
            [data-testid="stTextArea"] textarea,
            [data-testid="stTextInput"] input,
            [data-testid="stNumberInput"] input {
                background-color: #2D2D2D !important;
                color: #FFFFFF !important;
                border: 1px solid #555555 !important;
            }
            .stTextArea textarea:focus {
                border-color: #4A90E2 !important;
                box-shadow: 0 0 0 0.2rem rgba(74, 144, 226, 0.25) !important;
            }
        </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stButton>button {
        background-color: grey !important;
        color: white !important;
        transition: all 0.3s ease !important;
    }
    .stButton>button:hover { opacity: 0.8 !important; transform: scale(1.05) !important; }
    [data-testid="expander"] > div:first-child {
        background-color: grey !important;
        color: white !important;
        border-radius: 8px !important;
    }
    </style>
""", unsafe_allow_html=True)

# File handlers implementation
def read_docx(file_content):
    """Read text from DOCX files"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"DOCX read error: {str(e)}")
        return ""

def read_pdf(content):
    """Extract text from PDF files"""
    try:
        pdf = PdfReader(io.BytesIO(content))
        text = []
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return '\n'.join(text)
    except Exception as e:
        st.error(f"PDF read error: {str(e)}")
        return ""

def read_md(content):
    """Process Markdown files"""
    try:
        return content.decode('utf-8')
    except Exception as e:
        st.error(f"Markdown read error: {str(e)}")
        return ""

def read_pages(content):
    """Extract text from Apple Pages files"""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            # Check for PDF preview first
            for name in z.namelist():
                if name.endswith('.pdf'):
                    return read_pdf(z.read(name))
            
            # Fallback to XML parsing
            xml_content = z.read('Index.xml')
            return parse_pages_xml(xml_content)
    except Exception as e:
        st.error(f"Pages file error: {str(e)}")
        return ""

def parse_pages_xml(xml_content):
    """Parse Pages XML content"""
    try:
        namespace = {'ns': 'http://developer.apple.com/namespaces/bl'}
        root = ET.fromstring(xml_content)
        text_elements = root.findall('.//ns:t', namespace)
        return '\n'.join([elem.text for elem in text_elements if elem.text])
    except Exception as e:
        st.error(f"XML parse error: {str(e)}")
        return ""

# Sidebar configuration
with st.sidebar:
    st.header("Canvas Access Information")
    canvas_domain = st.text_input("Canvas Domain", value="elu.instructure.com", key="canvas_domain_input")
    course_id = st.text_input("Course ID", value="424", key="course_id_input")
    st.session_state.course_id = course_id
    st.markdown("<style> .css-1m8p7c5 { background-color: #f4f4f9; padding: 15px; border-radius: 8px; } .css-12ttj6a { background-color: #ff6347; color: white; padding: 10px 20px; border-radius: 10px; } .css-12ttj6a:hover { background-color: #e5533b; } .css-18e3th9 { border-radius: 12px; border: 2px solid #ff6347; padding: 10px; } </style>", unsafe_allow_html=True)
    st.sidebar.markdown("""<span style='color:red;'>Fetch only when database is empty!</span>""", unsafe_allow_html=True)

# User records and main dashboard
records_staff = {
    # "victor.lucas@elu.nl": {"name": "Victor Lucas", "student_id": "182"},
    # "victor@amsterdam.tech": {"name": "Victor Lucas", "student_id": "464"},
    # "test_instructor@amsterdam.tech": {"name": "Test_instructor", "student_id": "182"},
    # "test_student@amsterdam.tech": {"name": "Test_student", "student_id": "182"},
    # "hugo@elu.nl": {"name": "Hugo", "student_id": "182"},
    # "isadora@amsterdam.tech": {"name": "Isadora", "student_id": "464"},
    # "didem@amsterdam.tech": {"name": "Didem", "student_id": "464"},
    # "alper@amsterdam.tech": {"name": "Alper", "student_id": "182"},
    # "henry@amsterdam.tech": {"name": "Henry", "student_id": "464"},
    "moisieiev.mykyta@elu.nl": {"name": "Moisieiev Mykyta", "student_id": "139"},
    "erika.parra@elu.nl": {"name": "Erika Parra", "student_id": "464"}, 
    "leon.aziz@elu.nl": {"name": "Leon Aziz", "student_id": "182"},
    "yana.kondratyuk@elu.nl": {"name": "Yana Kondratyuk", "student_id": "165"},
    "jozsef.dudas@elu.nl": {"name": "Dudás József", "student_id": "403"},
    "viktor.rodau@elu.nl": {"name": "Viktor Rodau", "student_id": "181"},
    "isadora@amsterdam.tech": {"name": "Isadora Costa", "student_id": "2655"}
}

def get_user_details(email):
    record = records_staff.get(email)
    return (record["name"], record["student_id"]) if record else (None, None)

def instructor_dashboard():
    track_activity()
    initialize_session()
    restrict_access(allowed_roles=['instructor'])
    
    session_defaults = {
        'last_saved_record': None,
        'show_preview': False,
        'previous_student': None
    }
    for key, value in session_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    st.title("Instructor Dashboard")

    # Initialize database connection with robust error handling
    conn = None
    try:
        conn = get_db_connection()
        if conn is None:
            st.error("⚠️ Database connection failed. Please check your database configuration.")
            st.stop()
    except Exception as e:
        st.error(f"⚠️ Database connection error: {str(e)}")
        return

    if st.sidebar.button("Fetch Submissions"):
        fetch_submissions(canvas_domain, course_id)
        st.success("Submissions fetched successfully.")
        st.markdown("""<span style='color:red;'>Please, use this button only when database is empty</span>""", 
                    unsafe_allow_html=True)

    try:
        # Use context manager for database operations
        with conn:
            with conn.cursor() as cursor:
                st.markdown("<h5 style='font-size: 24px;'>Grade Submissions</h5>", unsafe_allow_html=True)

                # Course and assignment selection
                cursor.execute("SELECT DISTINCT course_id FROM submissions")
                db_courses = [row[0] for row in cursor.fetchall()]
                courses = ["Nil"] + db_courses
                selected_course = st.selectbox("Select Course", courses, index=0)

                if selected_course == "Nil":
                    assignments = []
                else:
                    cursor.execute("""
                        SELECT assignment_id, assignment_name 
                        FROM submissions 
                        WHERE course_id = %s
                        GROUP BY assignment_id, assignment_name
                    """, (selected_course,))
                    db_assignments = cursor.fetchall()
                    assignments = [("Nil", "Nil")] + db_assignments

                if not assignments:
                    st.warning("No assignments found for this course")
                    return

                assignment_options = {f"{a[0]} - {a[1]}": a for a in assignments}
                selected_assignment = st.selectbox("Select Assignment", list(assignment_options.keys()), index=0)
                assignment_id, assignment_name = assignment_options[selected_assignment]

                # Student selection
                if selected_course != "Nil" and assignment_id != "Nil":
                    cursor.execute("""
                        SELECT DISTINCT student_id 
                        FROM submissions 
                        WHERE course_id = %s AND assignment_id = %s
                    """, (selected_course, assignment_id))
                    db_students = [s[0] for s in cursor.fetchall()]
                    students = ["Nil"] + db_students
                else:
                    students = ["Nil"]

                selected_student = st.selectbox("Select Student", students, key="student_selector", index=0)

                # Student info display
                if selected_student != "Nil":
                    student_name = next((details["name"] for details in records_staff.values() 
                                        if str(details["student_id"]) == str(selected_student)), None)
                    st.markdown(f'<h5 style="color: red;">The document under review belongs to {student_name} and has the ID {selected_student}.</h5>', 
                                unsafe_allow_html=True)

                    if selected_student != st.session_state.previous_student:
                        keys_to_clear = [key for key in st.session_state.keys() 
                                        if key.startswith(('grade_', 'feedback_', 'suggestions_'))]
                        for key in keys_to_clear:
                            if key in st.session_state:
                                del st.session_state[key]
                        st.session_state.previous_student = selected_student

                # File handling with new formats
                MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
                if selected_course != "Nil" and assignment_id != "Nil" and selected_student != "Nil":
                    cursor.execute("""
                        SELECT DISTINCT file_path 
                        FROM submissions 
                        WHERE course_id = %s AND assignment_id = %s AND student_id = %s
                    """, (selected_course, assignment_id, selected_student))
                    files = list(set(cursor.fetchall()))
                else:
                    files = []

                if files:
                    displayed_files = set()
                    for idx, file_info in enumerate(files):
                        file_path = file_info[0]
                        if file_path in displayed_files:
                            continue
                        displayed_files.add(file_path)

                        text = ""
                        with st.expander(f"View {os.path.basename(file_path)}"):
                            blob_client = container_client.get_blob_client(file_path)
                            content = blob_client.download_blob().readall()

                            # File size check
                            if len(content) > MAX_FILE_SIZE:
                                st.error(f"File too large to process ({len(content)/1024/1024:.1f}MB)")
                                continue

                            formatted_key = f"formatted_{file_path}"
                            if formatted_key not in st.session_state:
                                st.session_state[formatted_key] = None

                            # File type handling
                            if file_path.endswith('.docx'):
                                text = read_docx(content)
                            elif file_path.endswith('.pdf'):
                                text = read_pdf(content)
                            elif file_path.endswith('.md'):
                                text = read_md(content)
                            elif file_path.endswith('.pages'):
                                text = read_pages(content)
                            elif file_path.endswith('.txt'):
                                text = content.decode('utf-8')
                            else:
                                text = f"Unsupported file format: {file_path.split('.')[-1]}"

                            # Formatting controls and UI
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                if st.button("🖋️ View Formatted Document", 
                                           key=f"format_{file_path}_{idx}",
                                           help="Enhance document formatting for better readability"):
                                    with st.spinner("Improving document formatting..."):
                                        st.session_state[formatted_key] = format_text_for_display(text)
                            
                            with col2:
                                if st.session_state[formatted_key]:
                                    if st.button("❌ Clear Formatting", 
                                               key=f"clear_{file_path}_{idx}"):
                                        st.session_state[formatted_key] = None

                            # Text display
                            if st.session_state[formatted_key]:
                                st.text_area("Formatted Content", 
                                           st.session_state[formatted_key], 
                                           height=400,
                                           key=f"formatted_area_{idx}")
                            else:
                                st.text_area("Document Content", 
                                           text, 
                                           height=400,
                                           key=f"raw_area_{idx}")

                            # Rubric and feedback section
                            rubric = fetch_assignment_rubric(assignment_name)
                            if st.button("View Rubric", key=f"view_rubric_{idx}"):
                                keys_to_clear = [
                                    f"grade_{idx}", 
                                    f"feedback_{idx}",
                                    f"suggestions_{idx}",
                                    f"instructor_comment_{idx}",
                                    f"grading_saved_{idx}"
                                ]
                                for key in keys_to_clear:
                                    if key in st.session_state:
                                        del st.session_state[key]

                                st.markdown(
                                    """<span style='color:red;'>Click the button at least two times, and each click will prompt the AI to refined rubric table!</span>""",
                                    unsafe_allow_html=True
                                )
                                with st.spinner("Converting rubric to table..."):
                                    rubric_table = convert_rubric_to_table(rubric)
                                if rubric_table:
                                    st.markdown("### Assignment Rubric")
                                    st.markdown(rubric_table, unsafe_allow_html=True)
                                else:
                                    st.error("Could not generate rubric table")

                            if st.button(f"Generate AI Feedback", key=f"ai_feedback_{idx}"):
                                with st.spinner("Generating feedback..."):
                                    result = get_ai_feedback(rubric, text)
                                    if "error" in result:
                                        st.error(f"AI Error: {result['error']}")
                                    else:
                                        st.session_state[f"grade_{idx}"] = result['grade']
                                        st.session_state[f"feedback_{idx}"] = result['feedback']
                                        st.session_state[f"suggestions_{idx}"] = result['suggestions']

                            # Feedback editing and saving
                            if f"feedback_{idx}" in st.session_state:
                                st.subheader("Review AI Grading")
                                
                                grade = st.text_area(
                                    "Edit Grade", 
                                    value=st.session_state[f"grade_{idx}"],
                                    key=f"grade_edit_{idx}",
                                    height=68
                                )
                                feedback = st.text_area(
                                    "Edit Feedback", 
                                    value=st.session_state[f"feedback_{idx}"],
                                    key=f"feedback_edit_{idx}",
                                    height=300
                                )
                                suggestions = st.text_area(
                                    "Edit Suggestions", 
                                    value=st.session_state[f"suggestions_{idx}"],
                                    key=f"suggestions_edit_{idx}",
                                    height=300
                                )
                                
                                st.markdown('**<span style="color: red;">Important!</span> Does this feedback reflect the assessment criteria?**', 
                                           unsafe_allow_html=True)
                                instructor_comment = st.text_area(
                                    "Instructor Comments",
                                    key=f"instructor_comment_{idx}",
                                    height=150
                                )
                                
                                if st.button("Approve & Save", key=f"approve_save_{idx}"):
                                    if not instructor_comment.strip():
                                        st.error("You must provide feedback before continuing.")
                                    else:
                                        st.session_state[f"grade_{idx}"] = grade
                                        st.session_state[f"feedback_{idx}"] = feedback
                                        st.session_state[f"suggestions_{idx}"] = suggestions

                                        record = {
                                            "course_id": selected_course,
                                            "assignment": assignment_name,
                                            "student": selected_student,
                                            "file": file_path,
                                            "grade": grade,
                                            "feedback": feedback,
                                            "suggestions": suggestions,
                                            "instructor_comment": instructor_comment,
                                            "timestamp": pd.Timestamp.now().isoformat()
                                        }

                                        save_result = save_grading_record(record)
                                        
                                        if save_result == 'success':
                                            st.session_state[f"grading_saved_{idx}"] = "success"
                                            st.session_state.last_saved_record = record
                                        elif save_result == 'duplicate':
                                            st.session_state[f"grading_saved_{idx}"] = "duplicate"
                                        else:
                                            st.session_state[f"grading_saved_{idx}"] = "error"

                            # Handle save status
                            if st.session_state.get(f"grading_saved_{idx}") == "success":
                                st.success("Grading saved successfully.")
                                st.markdown("""
                                    <script>   
                                    setTimeout(function() {    
                                        var elements = document.querySelectorAll('.stAlert');
                                        if (elements.length > 0) {     
                                            elements[elements.length - 1].style.display = 'none';
                                        }
                                    }, 2000);
                                    </script>        
                                """, unsafe_allow_html=True)            
                                st.session_state[f"grading_saved_{idx}"] = None
                                            
                            elif st.session_state.get(f"grading_saved_{idx}") == "duplicate":
                                st.error("Duplicate entry - already saved!")
                                st.markdown("""
                                    <script>   
                                    setTimeout(function() {    
                                        var elements = document.querySelectorAll('.stAlert');
                                        if (elements.length > 0) {     
                                            elements[elements.length - 1].style.display = 'none';
                                        }
                                    }, 2000);
                                    </script>        
                                """, unsafe_allow_html=True)            
                                st.session_state[f"grading_saved_{idx}"] = None
                                
                            elif st.session_state.get(f"grading_saved_{idx}") == "error":
                                st.error("Error saving feedback to database.")
    except Exception as e:
        st.error(f"Database operation failed: {str(e)}")
    finally:
        if conn:
            conn.close()

    # Export Section
    st.markdown("---")
    st.subheader("Export Grading Records")

    time_filter = st.selectbox(
        "Select Time Range",
        options=['last_hour', 'last_day', 'last_week', 'last_month', 'last_year', 'all_time'],
        format_func=lambda x: x.replace('_', ' ').title(),
        key='global_time_filter'
    )

    if st.button("Preview Filtered Records"):
        try:
            records_df = get_grading_records(time_filter)
            if not records_df.empty:
                st.dataframe(records_df)
            else:
                st.warning("No records found")
        except Exception as e:
            st.error(f"Failed to load records: {str(e)}")

    if st.button("Export to Excel"):
        try:
            records_df = get_grading_records(time_filter)
            if not records_df.empty:
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    records_df.to_excel(writer, index=False)
                st.download_button(
                    label="Download Excel File",
                    data=output.getvalue(),
                    file_name=f"grading_records_{time_filter}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            else:
                st.warning("No records to export")
        except Exception as e:
            st.error(f"Export failed: {str(e)}")

if __name__ == "__main__":
    instructor_dashboard()


# import streamlit as st
# st.set_page_config(initial_sidebar_state="expanded")
# from utils.session import initialize_session, restrict_access, update_activity
# from utils.canvas_integration import fetch_submissions
# from utils.grading import fetch_assignment_rubric, get_ai_feedback, save_grading_record, convert_rubric_to_table, format_text_for_display
# from database.operations import get_grading_records
# from database.connection import get_db_connection
# from utils.security import get_all_grading_data
# from utils.azure_storage import *
# import pandas as pd
# import docx
# import io
# import os
# import zipfile
# from PyPDF2 import PdfReader
# from xml.etree import ElementTree as ET


# import streamlit as st

# # Remove ONLY the "About" option from the menu
# hide_about_option = """
# <style>
# /* Hide the About menu item */
# li[data-testid="main-menu-about"] {
#     display: none !important;
# }
# </style>
# """
# st.markdown(hide_about_option, unsafe_allow_html=True)


# # Initialize session and security
# def track_activity():
#     if st.session_state.get("logged_in"):
#         update_activity()

# # Theme customization
# current_theme = st.get_option("theme.base")
# if current_theme == "dark":
#     st.markdown("""
#         <style>
#             [data-testid="stTextArea"] textarea,
#             [data-testid="stTextInput"] input,
#             [data-testid="stNumberInput"] input {
#                 background-color: #2D2D2D !important;
#                 color: #FFFFFF !important;
#                 border: 1px solid #555555 !important;
#             }
#             .stTextArea textarea:focus {
#                 border-color: #4A90E2 !important;
#                 box-shadow: 0 0 0 0.2rem rgba(74, 144, 226, 0.25) !important;
#             }
#         </style>
#     """, unsafe_allow_html=True)

# st.markdown("""
#     <style>
#     .stButton>button {
#         background-color: grey !important;
#         color: white !important;
#         transition: all 0.3s ease !important;
#     }
#     .stButton>button:hover { opacity: 0.8 !important; transform: scale(1.05) !important; }
#     [data-testid="expander"] > div:first-child {
#         background-color: grey !important;
#         color: white !important;
#         border-radius: 8px !important;
#     }
#     </style>
# """, unsafe_allow_html=True)

# # File handlers implementation
# def read_docx(file_content):
#     """Read text from DOCX files"""
#     try:
#         doc = docx.Document(io.BytesIO(file_content))
#         return '\n'.join([para.text for para in doc.paragraphs])
#     except Exception as e:
#         st.error(f"DOCX read error: {str(e)}")
#         return ""

# def read_pdf(content):
#     """Extract text from PDF files"""
#     try:
#         pdf = PdfReader(io.BytesIO(content))
#         text = []
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text.append(page_text)
#         return '\n'.join(text)
#     except Exception as e:
#         st.error(f"PDF read error: {str(e)}")
#         return ""

# def read_md(content):
#     """Process Markdown files"""
#     try:
#         return content.decode('utf-8')
#     except Exception as e:
#         st.error(f"Markdown read error: {str(e)}")
#         return ""

# def read_pages(content):
#     """Extract text from Apple Pages files"""
#     try:
#         with zipfile.ZipFile(io.BytesIO(content)) as z:
#             # Check for PDF preview first
#             for name in z.namelist():
#                 if name.endswith('.pdf'):
#                     return read_pdf(z.read(name))
            
#             # Fallback to XML parsing
#             xml_content = z.read('Index.xml')
#             return parse_pages_xml(xml_content)
#     except Exception as e:
#         st.error(f"Pages file error: {str(e)}")
#         return ""

# def parse_pages_xml(xml_content):
#     """Parse Pages XML content"""
#     try:
#         namespace = {'ns': 'http://developer.apple.com/namespaces/bl'}
#         root = ET.fromstring(xml_content)
#         text_elements = root.findall('.//ns:t', namespace)
#         return '\n'.join([elem.text for elem in text_elements if elem.text])
#     except Exception as e:
#         st.error(f"XML parse error: {str(e)}")
#         return ""

# # Sidebar configuration
# with st.sidebar:
#     st.header("Canvas Access Information")
#     canvas_domain = st.text_input("Canvas Domain", value="elu.instructure.com", key="canvas_domain_input")
#     course_id = st.text_input("Course ID", value="424", key="course_id_input")
#     st.session_state.course_id = course_id
#     st.markdown("<style> .css-1m8p7c5 { background-color: #f4f4f9; padding: 15px; border-radius: 8px; } .css-12ttj6a { background-color: #ff6347; color: white; padding: 10px 20px; border-radius: 10px; } .css-12ttj6a:hover { background-color: #e5533b; } .css-18e3th9 { border-radius: 12px; border: 2px solid #ff6347; padding: 10px; } </style>", unsafe_allow_html=True)
#     st.sidebar.markdown("""<span style='color:red;'>Fetch only when database is empty!</span>""", unsafe_allow_html=True)

# # User records and main dashboard
# records_staff = {
#     # "victor.lucas@elu.nl": {"name": "Victor Lucas", "student_id": "182"},
#     # "victor@amsterdam.tech": {"name": "Victor Lucas", "student_id": "464"},
#     # "test_instructor@amsterdam.tech": {"name": "Test_instructor", "student_id": "182"},
#     # "test_student@amsterdam.tech": {"name": "Test_student", "student_id": "182"},
#     # "hugo@elu.nl": {"name": "Hugo", "student_id": "182"},
#     # "isadora@amsterdam.tech": {"name": "Isadora", "student_id": "464"},
#     # "didem@amsterdam.tech": {"name": "Didem", "student_id": "464"},
#     # "alper@amsterdam.tech": {"name": "Alper", "student_id": "182"},
#     # "henry@amsterdam.tech": {"name": "Henry", "student_id": "464"},
#     "moisieiev.mykyta@elu.nl": {"name": "Moisieiev Mykyta", "student_id": "139"},
#     "erika.parra@elu.nl": {"name": "Erika Parra", "student_id": "464"}, 
#     "leon.aziz@elu.nl": {"name": "Leon Aziz", "student_id": "182"},
#     "yana.kondratyuk@elu.nl": {"name": "Yana Kondratyuk", "student_id": "165"},
#     "jozsef.dudas@elu.nl": {"name": "Dudás József", "student_id": "403"},
#     "viktor.rodau@elu.nl": {"name": "Viktor Rodau", "student_id": "181"},
#     "isadora@amsterdam.tech": {"name": "Isadora Costa", "student_id": "2655"}
# }

# def get_user_details(email):
#     record = records_staff.get(email)
#     return (record["name"], record["student_id"]) if record else (None, None)

# def instructor_dashboard():
#     track_activity()
#     initialize_session()
#     restrict_access(allowed_roles=['instructor'])
    
#     session_defaults = {
#         'last_saved_record': None,
#         'show_preview': False,
#         'previous_student': None
#     }
#     for key, value in session_defaults.items():
#         if key not in st.session_state:
#             st.session_state[key] = value

#     st.title("Instructor Dashboard")

#     if st.sidebar.button("Fetch Submissions"):
#         fetch_submissions(canvas_domain, course_id)
#         st.success("Submissions fetched successfully.")
#         st.markdown("""<span style='color:red;'>Please, use this button only when database is empty</span>""", 
#                     unsafe_allow_html=True)

#     conn = get_db_connection()
#     cursor = conn.cursor()
#     st.markdown("<h5 style='font-size: 24px;'>Grade Submissions</h5>", unsafe_allow_html=True)

#     # Course and assignment selection
#     cursor.execute("SELECT DISTINCT course_id FROM submissions")
#     db_courses = [row[0] for row in cursor.fetchall()]
#     courses = ["Nil"] + db_courses
#     selected_course = st.selectbox("Select Course", courses, index=0)

#     if selected_course == "Nil":
#         assignments = []
#     else:
#         cursor.execute("""
#             SELECT assignment_id, assignment_name 
#             FROM submissions 
#             WHERE course_id = %s
#             GROUP BY assignment_id, assignment_name
#         """, (selected_course,))
#         db_assignments = cursor.fetchall()
#         assignments = [("Nil", "Nil")] + db_assignments

#     if not assignments:
#         st.warning("No assignments found for this course")
#         return

#     assignment_options = {f"{a[0]} - {a[1]}": a for a in assignments}
#     selected_assignment = st.selectbox("Select Assignment", list(assignment_options.keys()), index=0)
#     assignment_id, assignment_name = assignment_options[selected_assignment]

#     # Student selection
#     if selected_course != "Nil" and assignment_id != "Nil":
#         cursor.execute("""
#             SELECT DISTINCT student_id 
#             FROM submissions 
#             WHERE course_id = %s AND assignment_id = %s
#         """, (selected_course, assignment_id))
#         db_students = [s[0] for s in cursor.fetchall()]
#         students = ["Nil"] + db_students
#     else:
#         students = ["Nil"]

#     selected_student = st.selectbox("Select Student", students, key="student_selector", index=0)

#     # Student info display
#     if selected_student != "Nil":
#         student_name = next((details["name"] for details in records_staff.values() 
#         # student_name = next((details["name"] for details in records_staff.values() 
#                             if str(details["student_id"]) == str(selected_student)), None)
#         st.markdown(f'<h5 style="color: red;">The document under review belongs to {student_name} and has the ID {selected_student}.</h5>', 
#                     unsafe_allow_html=True)

#         if selected_student != st.session_state.previous_student:
#             keys_to_clear = [key for key in st.session_state.keys() 
#                             if key.startswith(('grade_', 'feedback_', 'suggestions_'))]
#             for key in keys_to_clear:
#                 del st.session_state[key]
#             st.session_state.previous_student = selected_student

#     # File handling with new formats
#     MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
#     if selected_course != "Nil" and assignment_id != "Nil" and selected_student != "Nil":
#         cursor.execute("""
#             SELECT DISTINCT file_path 
#             FROM submissions 
#             WHERE course_id = %s AND assignment_id = %s AND student_id = %s
#         """, (selected_course, assignment_id, selected_student))
#         files = list(set(cursor.fetchall()))
#     else:
#         files = []

#     if files:
#         displayed_files = set()
#         for idx, file_info in enumerate(files):
#             file_path = file_info[0]
#             if file_path in displayed_files:
#                 continue
#             displayed_files.add(file_path)

#             text = ""
#             with st.expander(f"View {os.path.basename(file_path)}"):
#                 blob_client = container_client.get_blob_client(file_path)
#                 content = blob_client.download_blob().readall()

#                 # File size check
#                 if len(content) > MAX_FILE_SIZE:
#                     st.error(f"File too large to process ({len(content)/1024/1024:.1f}MB)")
#                     continue

#                 formatted_key = f"formatted_{file_path}"
#                 if formatted_key not in st.session_state:
#                     st.session_state[formatted_key] = None

#                 # File type handling
#                 if file_path.endswith('.docx'):
#                     text = read_docx(content)
#                 elif file_path.endswith('.pdf'):
#                     text = read_pdf(content)
#                 elif file_path.endswith('.md'):
#                     text = read_md(content)
#                 elif file_path.endswith('.pages'):
#                     text = read_pages(content)
#                 elif file_path.endswith('.txt'):
#                     text = content.decode('utf-8')
#                 else:
#                     text = f"Unsupported file format: {file_path.split('.')[-1]}"

#                 # Formatting controls and UI
#                 col1, col2 = st.columns([3, 1])
#                 with col1:
#                     if st.button("🖋️ View Formatted Document", 
#                                key=f"format_{file_path}_{idx}",
#                                help="Enhance document formatting for better readability"):
#                         with st.spinner("Improving document formatting..."):
#                             st.session_state[formatted_key] = format_text_for_display(text)
                
#                 with col2:
#                     if st.session_state[formatted_key]:
#                         if st.button("❌ Clear Formatting", 
#                                    key=f"clear_{file_path}_{idx}"):
#                             st.session_state[formatted_key] = None

#                 # Text display
#                 if st.session_state[formatted_key]:
#                     st.text_area("Formatted Content", 
#                                st.session_state[formatted_key], 
#                                height=400,
#                                key=f"formatted_area_{idx}")
#                 else:
#                     st.text_area("Document Content", 
#                                text, 
#                                height=400,
#                                key=f"raw_area_{idx}")

#                 # Rubric and feedback section
#                 rubric = fetch_assignment_rubric(assignment_name)
#                 if st.button("View Rubric", key=f"view_rubric_{idx}"):
#                     keys_to_clear = [
#                         f"grade_{idx}", 
#                         f"feedback_{idx}",
#                         f"suggestions_{idx}",
#                         f"instructor_comment_{idx}",
#                         f"grading_saved_{idx}"
#                     ]
#                     for key in keys_to_clear:
#                         if key in st.session_state:
#                             del st.session_state[key]

#                     st.markdown(
#                         """<span style='color:red;'>Click the button at least two times, and each click will prompt the AI to refined rubric table!</span>""",
#                         unsafe_allow_html=True
#                     )
#                     with st.spinner("Converting rubric to table..."):
#                         rubric_table = convert_rubric_to_table(rubric)
#                     if rubric_table:
#                         st.markdown("### Assignment Rubric")
#                         st.markdown(rubric_table, unsafe_allow_html=True)
#                     else:
#                         st.error("Could not generate rubric table")

#                 if st.button(f"Generate AI Feedback", key=f"ai_feedback_{idx}"):
#                     with st.spinner("Generating feedback..."):
#                         result = get_ai_feedback(rubric, text)
#                         if "error" in result:
#                             st.error(f"AI Error: {result['error']}")
#                         else:
#                             st.session_state[f"grade_{idx}"] = result['grade']
#                             st.session_state[f"feedback_{idx}"] = result['feedback']
#                             st.session_state[f"suggestions_{idx}"] = result['suggestions']

#                 # Feedback editing and saving
#                 if f"feedback_{idx}" in st.session_state:
#                     st.subheader("Review AI Grading")
                    
#                     grade = st.text_area(
#                         "Edit Grade", 
#                         value=st.session_state[f"grade_{idx}"],
#                         key=f"grade_edit_{idx}",
#                         height=68
#                     )
#                     feedback = st.text_area(
#                         "Edit Feedback", 
#                         value=st.session_state[f"feedback_{idx}"],
#                         key=f"feedback_edit_{idx}",
#                         height=300
#                     )
#                     suggestions = st.text_area(
#                         "Edit Suggestions", 
#                         value=st.session_state[f"suggestions_{idx}"],
#                         key=f"suggestions_edit_{idx}",
#                         height=300
#                     )
                    
#                     st.markdown('**<span style="color: red;">Important!</span> Does this feedback reflect the assessment criteria?**', 
#                                unsafe_allow_html=True)
#                     instructor_comment = st.text_area(
#                         "Instructor Comments",
#                         key=f"instructor_comment_{idx}",
#                         height=150
#                     )
                    
#                     if st.button("Approve & Save", key=f"approve_save_{idx}"):
#                         if not instructor_comment.strip():
#                             st.error("You must provide feedback before continuing.")
#                         else:
#                             st.session_state[f"grade_{idx}"] = grade
#                             st.session_state[f"feedback_{idx}"] = feedback
#                             st.session_state[f"suggestions_{idx}"] = suggestions

#                             record = {
#                                 "course_id": selected_course,
#                                 "assignment": assignment_name,
#                                 "student": selected_student,
#                                 "file": file_path,
#                                 "grade": grade,
#                                 "feedback": feedback,
#                                 "suggestions": suggestions,
#                                 "instructor_comment": instructor_comment,
#                                 "timestamp": pd.Timestamp.now().isoformat()
#                             }

#                             save_result = save_grading_record(record)
                            
#                             if save_result == 'success':
#                                 st.session_state[f"grading_saved_{idx}"] = "success"
#                                 st.session_state.last_saved_record = record
#                             elif save_result == 'duplicate':
#                                 st.session_state[f"grading_saved_{idx}"] = "duplicate"
#                             else:
#                                 st.session_state[f"grading_saved_{idx}"] = "error"

#                     # Handle save status
#                     if st.session_state.get(f"grading_saved_{idx}") == "success":
#                         st.success("Grading saved successfully.")
#                         st.markdown("""
#                             <script>   
#                             setTimeout(function() {    
#                                 var elements = document.querySelectorAll('.stAlert');
#                                 if (elements.length > 0) {     
#                                     elements[elements.length - 1].style.display = 'none';
#                                 }
#                             }, 2000);
#                             </script>        
#                         """, unsafe_allow_html=True)            
#                         st.session_state[f"grading_saved_{idx}"] = None
                                    
#                     elif st.session_state.get(f"grading_saved_{idx}") == "duplicate":
#                         st.error("Duplicate entry - already saved!")
#                         st.markdown("""
#                             <script>   
#                             setTimeout(function() {    
#                                 var elements = document.querySelectorAll('.stAlert');
#                                 if (elements.length > 0) {     
#                                     elements[elements.length - 1].style.display = 'none';
#                                 }
#                             }, 2000);
#                             </script>        
#                         """, unsafe_allow_html=True)            
#                         st.session_state[f"grading_saved_{idx}"] = None
                        
#                     elif st.session_state.get(f"grading_saved_{idx}") == "error":
#                         st.error("Error saving feedback to database.")

#     # Export Section
#     st.markdown("---")
#     st.subheader("Export Grading Records")

#     time_filter = st.selectbox(
#         "Select Time Range",
#         options=['last_hour', 'last_day', 'last_week', 'last_month', 'last_year', 'all_time'],
#         format_func=lambda x: x.replace('_', ' ').title(),
#         key='global_time_filter'
#     )

#     if st.button("Preview Filtered Records"):
#         records_df = get_grading_records(time_filter)
#         if not records_df.empty:
#             st.dataframe(records_df)
#         else:
#             st.warning("No records found")

#     if st.button("Export to Excel"):
#         records_df = get_grading_records(time_filter)
#         if not records_df.empty:
#             output = io.BytesIO()
#             with pd.ExcelWriter(output, engine='openpyxl') as writer:
#                 records_df.to_excel(writer, index=False)
#             st.download_button(
#                 label="Download Excel File",
#                 data=output.getvalue(),
#                 file_name=f"grading_records_{time_filter}.xlsx",
#                 mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
#             )
#         else:
#             st.warning("No records to export")

# if __name__ == "__main__":
#     instructor_dashboard()





